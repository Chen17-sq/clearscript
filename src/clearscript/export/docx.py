"""DOCX exporter.

Renders the cleaned markdown transcript into an A4 document with Microsoft
YaHei body / 14pt bold title and dash-style bullets (matching the conventions
of the original personal skill that this project descends from).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _bootstrap_numbering(doc) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p._element.getparent().remove(p._element)


def _inject_dash_numbering(doc) -> int:
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    abs_id = (max(existing_abs) + 1) if existing_abs else 100
    num_id = (max(existing_num) + 1) if existing_num else 100

    abs_xml = f"""<w:abstractNum w:abstractNumId="{abs_id}" xmlns:w="{_W_NS}">
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/><w:numFmt w:val="bullet"/>
            <w:lvlText w:val="-"/><w:lvlJc w:val="left"/>
            <w:pPr><w:ind w:left="420" w:hanging="280"/></w:pPr>
            <w:rPr><w:rFonts w:ascii="Microsoft YaHei" w:hAnsi="Microsoft YaHei" w:eastAsia="Microsoft YaHei"/></w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/><w:numFmt w:val="bullet"/>
            <w:lvlText w:val="-"/><w:lvlJc w:val="left"/>
            <w:pPr><w:ind w:left="840" w:hanging="280"/></w:pPr>
            <w:rPr><w:rFonts w:ascii="Microsoft YaHei" w:hAnsi="Microsoft YaHei" w:eastAsia="Microsoft YaHei"/></w:rPr>
        </w:lvl>
    </w:abstractNum>"""
    num_xml = f"""<w:num w:numId="{num_id}" xmlns:w="{_W_NS}">
        <w:abstractNumId w:val="{abs_id}"/>
    </w:num>"""
    numbering.append(parse_xml(abs_xml))
    numbering.append(parse_xml(num_xml))
    return num_id


def _set_para_numbering(para, num_id: int, level: int = 0) -> None:
    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    nId = OxmlElement("w:numId")
    nId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(nId)
    pPr.append(numPr)


def _yahei(run, size: int = 11, bold: bool = False) -> None:
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rFonts.set(qn(attr), "Microsoft YaHei")


_SPEAKER_RE = re.compile(r"^([A-Za-z一-龥][A-Za-z一-龥\s\.]{0,30}?)[:：]\s*$")
_BULLET_RE = re.compile(r"^(\s*)-\s+(.+)$")


def write_docx(markdown: str, output_path: Path, *, title: str | None = None) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_height, section.page_width = Cm(29.7), Cm(21)
    for margin_attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin_attr, Cm(2.2))

    _bootstrap_numbering(doc)
    num_id = _inject_dash_numbering(doc)

    if title:
        p = doc.add_paragraph()
        _yahei(p.add_run(title), size=14, bold=True)

    for line in markdown.splitlines():
        if not line.strip():
            continue

        stripped = line.strip()
        speaker_match = _SPEAKER_RE.match(stripped)
        if speaker_match:
            doc.add_paragraph()
            p = doc.add_paragraph()
            label = speaker_match.group(1).strip()
            colon = "：" if any("一" <= ch <= "鿿" for ch in label) else ":"
            _yahei(p.add_run(f"{label}{colon}"), size=11, bold=True)
            continue

        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            level = 1 if len(bullet_match.group(1)) >= 2 else 0
            p = doc.add_paragraph()
            _yahei(p.add_run(bullet_match.group(2)), size=11)
            _set_para_numbering(p, num_id, level=level)
            continue

        p = doc.add_paragraph()
        _yahei(p.add_run(stripped), size=11)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
